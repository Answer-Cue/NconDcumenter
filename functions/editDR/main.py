import streamlit as st
from io import BytesIO
from docx import Document

st.header("NHK放送コンテスト 書類作成補佐")

# 1. 参加者情報フォーム
with st.form("participant_form"):
    st.subheader("参加者情報")
    name = st.text_input("氏名")
    school = st.text_input("学校名")
    grade = st.selectbox("学年", ["中1", "中2", "中3", "高1", "高2", "高3"])
    
    st.subheader("作品情報")
    title = st.text_input("作品タイトル")
    category = st.selectbox("カテゴリ", ["ラジオドラマ", "動画作品", "ドキュメンタリー"])
    
    submitted = st.form_submit_button("書類作成")

if submitted:
    st.success("入力を受け取りました！書類を生成します…")
    
    # 2. Word 書類テンプレートを作成
    doc = Document()
    doc.add_heading("NHK放送コンテスト応募書類", level=0)
    doc.add_paragraph(f"氏名: {name}")
    doc.add_paragraph(f"学校名: {school}")
    doc.add_paragraph(f"学年: {grade}")
    doc.add_paragraph(f"作品タイトル: {title}")
    doc.add_paragraph(f"カテゴリ: {category}")
    
    # 3. ダウンロード用バイナリ作成
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    st.download_button(
        label="書類をダウンロード",
        data=file_stream,
        file_name=f"{name}_応募書類.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
